from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parents[1]
DOC_OUTPUT = ROOT / "output" / "doc" / "project_i_report.docx"
PDF_OUTPUT = ROOT / "output" / "pdf" / "project_i_report.pdf"


REPORT = [
    ("title", "Project I Report: Sentinel Signature-Based Virus Scanner"),
    ("meta", "Network Security Practices - Project I"),
    ("heading", "1. Overview"),
    (
        "body",
        "Sentinel is a functional signature-based virus scanner. Its goal is to scan a target "
        "directory, compare file content against known malware signatures, apply basic heuristic "
        "rules, and generate a security report. The scanner is designed for safe demonstrations: "
        "it only reads files and never executes any scanned file.",
    ),
    (
        "body",
        "The demonstration uses the EICAR test file, a standard harmless antivirus test string, "
        "hidden in a nested folder. Sentinel detects it through hash signatures and byte-pattern "
        "matching, and also flags a separate suspicious sample through heuristic scoring.",
    ),
    ("heading", "2. System Functions"),
    ("bullet", "Recursive scanning of a target file or directory."),
    ("bullet", "MD5 and SHA256 calculation for exact signature matching."),
    ("bullet", "Hex and byte-pattern matching for known malware fragments."),
    (
        "bullet",
        "Heuristic scoring for suspicious strings such as CreateRemoteThread, WriteProcessMemory, "
        "VirtualAlloc, and -EncodedCommand.",
    ),
    (
        "bullet",
        "JSON and text report generation with summary counts, findings, timestamps, skipped files, "
        "and errors.",
    ),
    ("heading", "3. Signature Database Design"),
    (
        "body",
        "The signature database is stored in JSON because it is human-readable, easy to version "
        "control, and supported by Python's standard library. It contains three major sections: "
        "hashes, patterns, and heuristics.",
    ),
    (
        "body",
        "Hash signatures are stored in dictionaries keyed by digest value. This is appropriate "
        "because the scanner calculates a file digest once and then needs a fast membership test. "
        "Average dictionary lookup time is O(1), which keeps exact-signature matching efficient.",
    ),
    (
        "body",
        "Pattern signatures are stored as bytes decoded from hexadecimal strings. Heuristic rules "
        "store a target string or byte sequence, a score, and a threat level. A file becomes "
        "suspicious when its total heuristic score reaches the configured threshold.",
    ),
    ("heading", "4. Scanning Engine"),
    (
        "body",
        "Sentinel recursively traverses the target directory and scans only regular files. Symlinks "
        "are skipped to avoid loops or unintended traversal outside the target tree. Missing targets "
        "and read errors are recorded in the report instead of crashing the scanner.",
    ),
    (
        "body",
        "Files are read in chunks. Each chunk updates the MD5 and SHA256 hash contexts and is also "
        "searched for pattern and heuristic matches. The scanner keeps carry-over bytes between "
        "chunks so signatures split across chunk boundaries can still be detected.",
    ),
    ("heading", "5. Demonstration"),
    (
        "body",
        "The demo folder contains one clean file, one nested EICAR sample, and one suspicious mock "
        "file. The command used for the demo is:",
    ),
    (
        "code",
        "python3 -m sentinel scan --target samples --signatures signatures.json "
        "--output reports/report.json --text-output reports/report.txt",
    ),
    (
        "body",
        "The expected result is 3 scanned files, 1 clean file, 1 infected file, 1 suspicious file, "
        "0 skipped files, and 0 errors. The EICAR file is matched by SHA256, MD5, and byte pattern. "
        "The suspicious file reaches heuristic score 8 with a threshold of 5.",
    ),
    ("heading", "6. Testing"),
    (
        "body",
        "The test suite uses Python unittest and currently contains 9 tests. The tests cover "
        "signature database loading, invalid database handling, SHA256 and MD5 detection, hex "
        "pattern detection across chunk boundaries, heuristic scoring, empty folders, nested "
        "folders, large-file skipping, and missing target handling.",
    ),
    ("code", "python3 -m unittest discover -s tests"),
    ("heading", "7. System Features"),
    ("bullet", "Read-only scanning: scanned files are never executed."),
    ("bullet", "Efficient exact matching: hash signatures use dictionary lookup."),
    ("bullet", "Memory-aware scanning: files are processed in chunks."),
    (
        "bullet",
        "Explainable heuristic analysis: every suspicious finding includes matched rules and scores.",
    ),
    ("bullet", "Reproducible reports: the README command regenerates the JSON and text reports."),
    ("heading", "8. Limitations And Future Work"),
    (
        "body",
        "Sentinel is a teaching project, not a production antivirus engine. It does not unpack "
        "archives, emulate code execution, detect polymorphic malware, or deeply parse executable "
        "formats. A Bloom Filter could be added as a memory-efficient pre-check for very large "
        "signature databases, but exact hash-map confirmation would still be required to avoid "
        "false positives.",
    ),
    ("heading", "9. Conclusion"),
    (
        "body",
        "Sentinel satisfies Project I by implementing a safe, functional, and explainable "
        "signature-based virus scanner. It demonstrates known-signature detection, heuristic "
        "suspicious-file analysis, and comprehensive report generation with reproducible tests and "
        "demo output.",
    ),
]


def build_docx() -> None:
    DOC_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(15)

    for kind, text in REPORT:
        if kind == "title":
            paragraph = document.add_heading(text, level=0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif kind == "meta":
            paragraph = document.add_paragraph(text)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif kind == "heading":
            document.add_heading(text, level=1)
        elif kind == "bullet":
            document.add_paragraph(text, style="List Bullet")
        elif kind == "code":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            document.add_paragraph(text)

    document.save(DOC_OUTPUT)


def build_pdf() -> None:
    PDF_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ProjectTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        alignment=1,
        spaceAfter=12,
    )
    heading_style = ParagraphStyle(
        "ProjectHeading",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        spaceBefore=12,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "ProjectBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10,
        leading=13,
        spaceAfter=6,
    )
    code_style = ParagraphStyle(
        "ProjectCode",
        parent=styles["Code"],
        fontName="Courier",
        fontSize=8.5,
        leading=11,
        leftIndent=12,
        spaceAfter=8,
    )
    bullet_style = ParagraphStyle("ProjectBullet", parent=body_style, leftIndent=18, firstLineIndent=-10)

    story = []
    for kind, text in REPORT:
        safe_text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if kind == "title":
            story.append(Paragraph(safe_text, title_style))
        elif kind == "meta":
            story.append(Paragraph(safe_text, body_style))
            story.append(Spacer(1, 0.15 * inch))
        elif kind == "heading":
            story.append(Paragraph(safe_text, heading_style))
        elif kind == "bullet":
            story.append(Paragraph(f"- {safe_text}", bullet_style))
        elif kind == "code":
            story.append(Paragraph(safe_text, code_style))
        else:
            story.append(Paragraph(safe_text, body_style))

    doc = SimpleDocTemplate(
        str(PDF_OUTPUT),
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.7 * inch,
        bottomMargin=0.7 * inch,
    )
    doc.build(story)


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print(DOC_OUTPUT)
    print(PDF_OUTPUT)
